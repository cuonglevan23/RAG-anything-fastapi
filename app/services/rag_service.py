import os
import uuid
import asyncio
from functools import partial
from typing import Dict, Optional, Any
from pathlib import Path
from loguru import logger

from app.services.vlm_parser import CustomOpenAIPipeline
from raganything import RAGAnything, RAGAnythingConfig
from lightrag.llm.openai import openai_complete_if_cache, openai_embed
from lightrag.utils import EmbeddingFunc
from lightrag.rerank import cohere_rerank  # kept for optional fallback
from app.services.local_reranker import local_bge_rerank
# set_default_workspace: đảm bảo LightRAG dùng đúng namespace cho từng project
from lightrag.kg.shared_storage import set_default_workspace
from app.core.config import settings
from app.models.schema import ProcessingStatus
# test comment for commit

def _extract_docx_to_md(file_path: str, output_dir: str, file_basename: str) -> str:
    """
    Trích xuất text từ DOCX dùng python-docx high-level API.
    Cài đặt: pip install python-docx
    """
    try:
        from docx import Document
    except ImportError:
        raise ImportError("Cần cài python-docx: pip install python-docx")

    doc   = Document(file_path)
    lines = []

    # --- xuất tết cả paragraphs theo đúng thứ tự trong document ---
    # Ta dùng doc.paragraphs + doc.tables nhưng để giữ đúng thứ tự
    # (paragraph xen kẽ bảng) ta iterate doc.element.body như cũ nhưng
    # lấy nội dung qua python-docx objects

    para_idx  = 0
    table_idx = 0

    for child in doc.element.body:
        tag = child.tag.split("}")[-1]

        if tag == "p":
            if para_idx < len(doc.paragraphs):
                para  = doc.paragraphs[para_idx]
                para_idx += 1
                text  = para.text  # FULL text, kể cả runs
                style = para.style.name if para.style else ""

                if not text.strip():
                    lines.append("")
                    continue

                if style.startswith("Heading"):
                    try:
                        lvl = int(style.split()[-1])
                    except (ValueError, IndexError):
                        lvl = 2
                    lines.append(f"{'#' * min(lvl, 6)} {text.strip()}")
                elif style == "Title":
                    lines.append(f"# {text.strip()}")
                else:
                    lines.append(text.strip())

        elif tag == "tbl":
            if table_idx < len(doc.tables):
                tbl = doc.tables[table_idx]
                table_idx += 1
                try:
                    if tbl.rows:
                        header = [c.text.strip().replace("\n", " ") for c in tbl.rows[0].cells]
                        lines.append("| " + " | ".join(header) + " |")
                        lines.append("| " + " | ".join(["---"] * len(header)) + " |")
                        for row in tbl.rows[1:]:
                            lines.append("| " + " | ".join(
                                c.text.strip().replace("\n", " ") for c in row.cells) + " |")
                        lines.append("")
                except Exception as te:
                    logger.warning(f"Table extraction warning: {te}")

    # Loại bỏ nhiều dòng trống liên tiếp
    md_lines = []
    prev_empty = False
    for line in lines:
        is_empty = (line == "")
        if is_empty and prev_empty:
            continue
        md_lines.append(line)
        prev_empty = is_empty

    md_content = "\n".join(md_lines).strip()

    out_path = os.path.join(output_dir, file_basename, "vlm", f"{file_basename}.md")
    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    with open(out_path, "w", encoding="utf-8") as f:
        f.write(md_content)

    logger.info(f"DOCX extracted → {out_path} ({len(md_lines)} lines, {len(md_content)} chars)")
    return out_path


# Semaphore giới hạn chỉ 1 VLM model chạy tại một thời điểm → tránh CUDA OOM
# DOCX/TXT được xuất nằm ngoài semaphore → vẫn chạy song song
VLM_SEMAPHORE: asyncio.Semaphore = asyncio.Semaphore(1)



class RAGService:

    def __init__(self):
        self.instances: Dict[str, RAGAnything] = {}
        self.tasks: Dict[str, ProcessingStatus] = {}
        self._lock = asyncio.Lock()

    async def get_instance(self, project_id: str) -> RAGAnything:
        """Get or initialize RAGAnything instance for a project"""
        async with self._lock:
            if project_id in self.instances:
                return self.instances[project_id]

            logger.info(f"Initializing RAGAnything instance for project: {project_id}")
            
            project_dir = os.path.abspath(settings.BASE_RAG_DIR / project_id)
            os.makedirs(project_dir, exist_ok=True)
            
            # ✅ Scope LightRAG shared storage namespaces to THIS project
            # Without this, all projects share the same pipeline_status / namespace
            # causing cross-project data leakage when querying.
            set_default_workspace(project_id)
            logger.info(f"📰 Default workspace set to: {project_id}")
            
            logger.info(f"📁 PROJECT STORAGE PATH: {project_dir}")
            
            config = RAGAnythingConfig(
                working_dir=project_dir,
            )
            # ================================================================
            # FIX 1: Legal-domain system_prompt
            # Inject vào MỌI lần LLM được gọi (entity extraction, summarization,
            # query generation). Không set → LLM không biết ngữ cảnh pháp lý
            # → extract entity generic, thiếu nhận diện Chương/Mục/Điều.
            # ================================================================
            LEGAL_SYSTEM_PROMPT = """Bạn là chuyên gia phân tích văn bản pháp luật Việt Nam.
Nhiệm vụ: xử lý luật, nghị định, thông tư, quyết định của Nhà nước Việt Nam.

Quy tắc bắt buộc:
1. Nhận diện cấu trúc phân cấp: PHẦN > CHƯƠNG > MỤC > ĐIỀU > KHOẢN > ĐIỂM
2. Mỗi "Điều X" là đơn vị pháp lý độc lập — TRÍCH DẪN NGUYÊN VĂN, không tóm tắt
3. Quan hệ tham chiếu giữa các điều luật rất quan trọng (Điều A dẫn chiếu Điều B)
4. Giữ nguyên số điều, khoản, điểm chính xác (Điều 12 khoản 2 điểm a)
5. Trả lời bằng tiếng Việt, dùng thuật ngữ pháp lý chính xác"""

            def _sanitize(text: str) -> str:
                """Loại bỏ null bytes và ký tự UTF-8 không hợp lệ gây ra lỗi 400 khi gửi OpenAI."""
                # 1. Re-encode để loại bỏ byte không hợp lệ
                text = text.encode("utf-8", errors="ignore").decode("utf-8")
                # 2. Xóa null bytes và control chars (trừ \n, \t)
                text = "".join(ch for ch in text if ch == "\n" or ch == "\t" or (ord(ch) >= 32 and ch != "\x7f"))
                return text

            def llm_model_func(prompt, system_prompt=None, history_messages=[], **kwargs):
                # Sanitize prompt trước khi gửi → tránh 400 'invalid JSON body'
                prompt = _sanitize(prompt) if isinstance(prompt, str) else prompt
                effective_system = system_prompt if system_prompt is not None else LEGAL_SYSTEM_PROMPT
                return openai_complete_if_cache(
                    settings.LLM_MODEL, prompt, system_prompt=effective_system,
                    history_messages=history_messages, api_key=settings.OPENAI_API_KEY, **kwargs,
                )

            embedding_func = EmbeddingFunc(
                embedding_dim=1536, max_token_size=8192,
                func=lambda texts: openai_embed(texts, model=settings.EMBEDDING_MODEL, api_key=settings.OPENAI_API_KEY),
            )

            # ================================================================
            # FIX 2: Vietnamese Legal Chunker
            # Ưu tiên split tại ranh giới Chương/Mục/Điều TRƯỚC khi split token.
            # → Mỗi "Điều X" luôn nằm hoàn chỉnh trong ≥1 chunk.
            # → Giải quyết vấn đề điều luật cuối tài liệu bị mất nội dung.
            # ================================================================
            import re as _re

            def vietnamese_legal_chunker(
                tokenizer, content: str,
                split_by_character=None, split_by_character_only: bool = False,
                chunk_overlap_token_size: int = 200, chunk_token_size: int = 400,
            ):
                LEGAL_BOUNDARY = _re.compile(
                    r'(?=\n(?:PHẦN|CHƯƠNG|MỤC|ĐIỀU|Phần|Chương|Mục|Điều)\s+[\dIVXivx]+[\.:]?\s)',
                    _re.UNICODE,
                )
                parts = LEGAL_BOUNDARY.split(content)
                results = []
                chunk_idx = 0
                for part in parts:
                    if not part.strip():
                        continue
                    part = _sanitize(part)  # ← strip null bytes trước khi encode
                    tokens = tokenizer.encode(part)
                    if len(tokens) <= chunk_token_size:
                        results.append({"tokens": len(tokens), "content": part.strip(), "chunk_order_index": chunk_idx})
                        chunk_idx += 1
                    else:
                        for start in range(0, len(tokens), chunk_token_size - chunk_overlap_token_size):
                            sub = tokens[start: start + chunk_token_size]
                            results.append({"tokens": len(sub), "content": tokenizer.decode(sub).strip(), "chunk_order_index": chunk_idx})
                            chunk_idx += 1
                            if start + chunk_token_size >= len(tokens):
                                break
                if not results:  # Fallback: không tìm thấy boundary pháp lý
                    tokens = tokenizer.encode(content)
                    for i, start in enumerate(range(0, len(tokens), chunk_token_size - chunk_overlap_token_size)):
                        sub = tokens[start: start + chunk_token_size]
                        results.append({"tokens": len(sub), "content": tokenizer.decode(sub).strip(), "chunk_order_index": i})
                return results

            # ============================================================
            # Cohere Rerank 3.5 Configuration
            # Để bật: set RERANK_ENABLE=true và COHERE_API_KEY trong .env
            # ============================================================
            # Local BGE Reranker (free, ~560MB VRAM, multilingual)
            # Bật bằng: RERANK_ENABLE=true trong .env
            # Model mặc định: BAAI/bge-reranker-v2-m3
            # Cài đặt:  pip install FlagEmbedding
            # ============================================================
            rerank_func = None
            if settings.RERANK_ENABLE:
                rerank_func = partial(
                    local_bge_rerank,
                    model_name=settings.RERANK_MODEL,  # BAAI/bge-reranker-v2-m3
                )
                logger.info(f"✅ Local BGE Reranker enabled: {settings.RERANK_MODEL}")
            else:
                logger.info("⚠️  Reranker disabled. Set RERANK_ENABLE=true in .env to enable.")
            # ============================================================

            # ================================================================
            # FIX 3: Legal entity_types — thay thế DEFAULT generic types
            # DEFAULT = ["organization","person","geo","event"] hoàn toàn
            # không phù hợp với văn bản luật Việt Nam.
            # ================================================================
            lightrag_kwargs = {
                "workspace": project_id,
                "chunk_token_size": 400,
                "chunk_overlap_token_size": 200,
                "chunking_func": vietnamese_legal_chunker,  # FIX 2
                "addon_params": {
                    "language": "Vietnamese",
                    "entity_extract_max_gleaning": 2,
                    "insert_batch_size": 5,
                    "entity_types": [       # FIX 3
                        "dieu",             # Điều X — đơn vị pháp lý cơ bản
                        "khoan",            # Khoản trong Điều
                        "muc",              # Mục trong Chương
                        "chuong",           # Chương trong Luật
                        "phan",             # Phần (cấp cao nhất)
                        "to_chuc",          # Tổ chức, cơ quan nhà nước
                        "khai_niem",        # Định nghĩa, khái niệm pháp lý
                        "hanh_vi",          # Hành vi, hoạt động được quy định
                        "doi_tuong",        # Đối tượng áp dụng
                        "chinh_sach",       # Chính sách, quy định chung
                    ],
                },
            }
            if rerank_func:
                lightrag_kwargs["rerank_model_func"] = rerank_func  # Bật Cohere Reranker

            rag = RAGAnything(
                config=config,
                llm_model_func=llm_model_func,
                embedding_func=embedding_func,
                lightrag_kwargs=lightrag_kwargs,
            )
            
            await rag._ensure_lightrag_initialized()
            self.instances[project_id] = rag
            logger.info(f"Project '{project_id}' initialized successfully.")
            return rag

    async def list_projects(self) -> list[str]:
        """List all available project IDs based on storage directories"""
        if not settings.BASE_RAG_DIR.exists():
            return []
        return [d.name for d in settings.BASE_RAG_DIR.iterdir() if d.is_dir()]

    async def finalize(self):
        """Cleanup all instances on shutdown"""
        for project_id, rag in self.instances.items():
            await rag.finalize_storages()
            logger.info(f"Project '{project_id}' finalized.")

    async def process_document(self, task_id: str, file_path: str, filename: str, project_id: str):
        """Background task for processing document in a specific project"""
        try:
            rag = await self.get_instance(project_id)
            
            # Step 1: Parsing
            async with self._lock:
                task = self.tasks[task_id]
                task.status = "parsing"
                task.message = "Parsing document structure..."
                task.percentage = 10.0
                task.logs.append(f"Starting engine: Mineru Parsing for project {project_id}...")
            # Step 1.1: Parse document using VLM (GPT-4o)
            task.logs.append(f"Starting engine: OpenAIPipeline (VLM) for project {project_id}...")
            
            # Khởi tạo pipeline và process
            vlm_pipeline = CustomOpenAIPipeline(api_key=settings.OPENAI_API_KEY)

            file_basename = os.path.splitext(os.path.basename(file_path))[0]
            project_dir   = os.path.abspath(settings.BASE_RAG_DIR / project_id)
            output_dir    = project_dir
            file_ext      = os.path.splitext(file_path)[1].lower()

            parsed_md_path = os.path.join(output_dir, file_basename, "vlm", f"{file_basename}.md")

            if os.path.exists(parsed_md_path):
                task.logs.append(f"Found existing Markdown at {parsed_md_path}. Skipping parsing.")

            elif file_ext in (".docx", ".doc"):
                # ── DOCX/DOC: convert sang PDF bằng LibreOffice → feed VLM pipeline ──
                task.logs.append(f"DOCX detected — converting to PDF via LibreOffice...")
                import asyncio, subprocess, shutil as _shutil

                def _convert_docx_to_pdf(src: str, out_dir: str) -> str:
                    """LibreOffice headless convert → trả về path của PDF vừa tạo."""
                    result = subprocess.run(
                        ["libreoffice", "--headless", "--convert-to", "pdf",
                         "--outdir", out_dir, src],
                        capture_output=True, text=True, timeout=120
                    )
                    if result.returncode != 0:
                        raise RuntimeError(
                            f"LibreOffice conversion failed: {result.stderr.strip()}"
                        )
                    # LibreOffice tạo file cùng tên nhưng đuôi .pdf
                    basename = os.path.splitext(os.path.basename(src))[0]
                    pdf_path = os.path.join(out_dir, basename + ".pdf")
                    if not os.path.exists(pdf_path):
                        raise FileNotFoundError(f"Converted PDF not found: {pdf_path}")
                    return pdf_path

                # temp dir để tránh conflict khi nhiều file cùng tên
                tmp_dir = os.path.join(output_dir, file_basename, "tmp_pdf")
                os.makedirs(tmp_dir, exist_ok=True)

                converted_pdf = await asyncio.to_thread(
                    _convert_docx_to_pdf, file_path, tmp_dir
                )
                task.logs.append(f"Converted to PDF: {converted_pdf} — now processing via VLM...")

                # Feed PDF vào VLM pipeline y hệt file PDF thường
                async with VLM_SEMAPHORE:
                    task.logs.append("PDF GPU slot acquired — parsing converted DOCX...")
                    parsed_md_path = await asyncio.to_thread(
                        vlm_pipeline.process_pdf, converted_pdf, output_dir, file_basename
                    )
                task.logs.append(f"DOCX→PDF→VLM done: {parsed_md_path}")


            elif file_ext in (".txt", ".md"):
                # ── TXT/MD: đọc thẳng, không cần bất kỳ parser nào ──
                task.logs.append(f"Plain text detected — reading directly...")
                os.makedirs(os.path.dirname(parsed_md_path), exist_ok=True)
                import shutil
                shutil.copy2(file_path, parsed_md_path)
                task.logs.append(f"Plain text copied to: {parsed_md_path}")

            elif file_ext in (".png", ".jpg", ".jpeg", ".webp"):
                # ── Image: dùng VLM để mô tả ──
                task.logs.append(f"Image detected — waiting for GPU slot...")
                vlm_pipeline = CustomOpenAIPipeline(api_key=settings.OPENAI_API_KEY)
                import asyncio
                async with VLM_SEMAPHORE:   # ← chỉ 1 GPU task cùng lúc
                    task.logs.append(f"Image GPU slot acquired — processing via VLM...")
                    parsed_md_path = await asyncio.to_thread(
                        vlm_pipeline.process_pdf, file_path, output_dir, file_basename
                    )
                task.logs.append(f"VLM Parsing finished: {parsed_md_path}")

            else:
                # ── PDF và các định dạng khác: dùng VLM pipeline ──
                task.logs.append(f"PDF detected — waiting for GPU slot...")
                vlm_pipeline = CustomOpenAIPipeline(api_key=settings.OPENAI_API_KEY)
                import asyncio
                try:
                    async with VLM_SEMAPHORE:   # ← chỉ 1 GPU task cùng lúc
                        task.logs.append(f"PDF GPU slot acquired — parsing...")
                        parsed_md_path = await asyncio.to_thread(
                            vlm_pipeline.process_pdf, file_path, output_dir, file_basename
                        )
                    task.logs.append(f"VLM Parsing finished. Markdown created at: {parsed_md_path}")
                except Exception as pdf_err:
                    if "PDFium" in str(pdf_err) or "Data format" in str(pdf_err):
                        # File bị corrupt hoặc thực ra là DOCX đổi đuôi → thử fallback
                        task.logs.append(
                            f"⚠️ PDFium failed ({pdf_err}). Trying python-docx fallback..."
                        )
                        try:
                            parsed_md_path = await asyncio.to_thread(
                                _extract_docx_to_md, file_path, output_dir, file_basename
                            )
                            task.logs.append(f"✅ Fallback DOCX extraction succeeded: {parsed_md_path}")
                        except Exception as docx_err:
                            raise RuntimeError(
                                f"File '{filename}' cannot be parsed as PDF or DOCX. "
                                f"PDF error: {pdf_err}. DOCX error: {docx_err}"
                            )
                    else:
                        raise



            async with self._lock:
                task.percentage = 40.0
                task.status = "indexing"
                task.message = "Indexing markdown content into LightRAG..."

            # Step 2: Index Text Content Directly
            if os.path.exists(parsed_md_path):
                with open(parsed_md_path, "r", encoding="utf-8") as f:
                    md_text = f.read()

                from lightrag.utils import compute_mdhash_id
                doc_id = compute_mdhash_id(md_text, prefix="doc-")
                file_name = os.path.basename(file_path)

                await rag.lightrag.ainsert(
                    input=md_text,
                    file_paths=file_name,
                    ids=doc_id
                )
                
                async with self._lock:
                    task.percentage = 70.0
                    task.logs.append("Markdown content successfully indexed into Knowledge Base.")
            else:
                raise FileNotFoundError(f"Markdown file not found after processing: {parsed_md_path}")
            
            # Note: Multimodal blocks (Tables/Equations) are already converted to text 
            # (HTML/LaTeX) within the markdown by VLM, so we don't need the legacy 
            # separate multimodal extraction step from Mineru outputs here.
            
            async with self._lock:
                task = self.tasks[task_id]
                task.status = "completed"
                task.message = "Processing completed successfully."
                task.percentage = 100.0
                task.logs.append("Finalized indexing. System ready.")
                
        except Exception as e:
            logger.error(f"Error processing document {filename} in {project_id}: {e}")
            async with self._lock:
                task = self.tasks[task_id]
                task.status = "failed"
                task.message = "Processing failed."
                task.logs.append(f"ERROR: {str(e)}")
                task.error = str(e)

    def create_task(self, filename: str, project_id: str) -> str:
        """Create a new task and return its ID"""
        task_id = str(uuid.uuid4())
        self.tasks[task_id] = ProcessingStatus(
            task_id=task_id,
            filename=filename,
            project_id=project_id,
            status="pending",
            message="Queued for processing",
            percentage=0.0
        )
        return task_id

    def get_task_status(self, task_id: str) -> Optional[ProcessingStatus]:
        """Get status of a specific task"""
        return self.tasks.get(task_id)

    async def query(self, project_id: str, query: str, mode: str = "hybrid",
                    top_k: int = 100, response_type: str = "Structured List") -> str:
        """Execute RAG query within a specific project"""
        set_default_workspace(project_id)
        rag = await self.get_instance(project_id)
        
        effective_mode = mode
        if settings.RERANK_ENABLE and mode == "hybrid":
            effective_mode = "mix"
            logger.info(f"Reranker active: auto-upgrade mode hybrid → mix")

        return await rag.aquery(
            query=query,
            mode=effective_mode,
            top_k=top_k,
            response_type=response_type,
        )

    async def query_with_context(self, project_id: str, query: str, mode: str = "hybrid",
                                  top_k: int = 100, response_type: str = "Structured List") -> dict:
        """
        Execute RAG query and return BOTH the final answer and the retrieved raw contexts.
        Used by the RAGAS evaluation tab.
        """
        # Switch default workspace TRƯỚC khi query để tránh đọc nhầm project khác
        set_default_workspace(project_id)
        rag = await self.get_instance(project_id)

        effective_mode = mode
        if settings.RERANK_ENABLE and mode == "hybrid":
            effective_mode = "mix"

        # Run the main answer query
        answer = await rag.aquery(
            query=query,
            mode=effective_mode,
            top_k=top_k,
            response_type=response_type,
        )

        # Run naive query to get raw retrieved passages (contexts for RAGAS)
        # Naive mode bypasses Knowledge Graph and returns raw chunks directly
        try:
            raw_context = await rag.aquery(
                query=query,
                mode="naive",
                top_k=20,               # Fewer chunks, focus on most relevant
                response_type="Multiple Paragraphs"
            )
            # Split into individual passages for RAGAS (expects List[str])
            contexts = [p.strip() for p in raw_context.split("\n\n") if p.strip()]
            if not contexts:
                contexts = [raw_context]  # fallback: treat whole response as one context
        except Exception as e:
            logger.warning(f"Naive context retrieval failed: {e}. Using answer as context fallback.")
            contexts = [answer]

        return {
            "query":    query,
            "answer":   answer,
            "contexts": contexts,
        }

rag_service = RAGService()
